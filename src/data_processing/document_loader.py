"""
Document loader for processing Ayurvedic texts and documents.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from PyPDF2 import PdfReader
from docx import Document
from src.config.settings import settings

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load and process documents from various formats."""
    
    def __init__(self):
        self.supported_formats = settings.supported_formats
        self.max_file_size_mb = settings.max_file_size_mb
    
    def load_documents_from_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of document dictionaries with metadata and content
        """
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.warning(f"Directory {directory_path} does not exist")
            return documents
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and self._is_supported_format(file_path):
                try:
                    doc = self.load_document(file_path)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    logger.error(f"Error loading document {file_path}: {e}")
        
        logger.info(f"Loaded {len(documents)} documents from {directory_path}")
        return documents
    
    def load_document(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """
        Load a single document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document dictionary with metadata and content, or None if failed
        """
        try:
            # Check file size
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                logger.warning(f"File {file_path} is too large ({file_size_mb:.2f}MB)")
                return None
            
            # Load based on file extension
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return self._load_pdf(file_path)
            elif extension == '.docx':
                return self._load_docx(file_path)
            elif extension == '.txt':
                return self._load_txt(file_path)
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return None
    
    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Load PDF document."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text_content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                return {
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "pdf",
                    "content": text_content.strip(),
                    "metadata": {
                        "num_pages": len(pdf_reader.pages),
                        "title": file_path.stem
                    }
                }
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
    
    def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """Load DOCX document."""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx",
                "content": text_content.strip(),
                "metadata": {
                    "num_paragraphs": len(doc.paragraphs),
                    "title": file_path.stem
                }
            }
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def _load_txt(self, file_path: Path) -> Dict[str, Any]:
        """Load TXT document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_type": "txt",
                "content": content.strip(),
                "metadata": {
                    "title": file_path.stem
                }
            }
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            raise
    
    def _is_supported_format(self, file_path: Path) -> bool:
        """Check if file format is supported."""
        return file_path.suffix.lower() in [f'.{fmt}' for fmt in self.supported_formats]
    
    def get_document_statistics(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Get statistics about loaded documents.
        
        Args:
            documents: List of loaded documents
            
        Returns:
            Dictionary with document statistics
        """
        if not documents:
            return {"total_documents": 0, "total_content_length": 0}
        
        total_content_length = sum(len(doc["content"]) for doc in documents)
        file_types = {}
        
        for doc in documents:
            file_type = doc["file_type"]
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            "total_documents": len(documents),
            "total_content_length": total_content_length,
            "file_types": file_types,
            "average_content_length": total_content_length / len(documents)
        } 